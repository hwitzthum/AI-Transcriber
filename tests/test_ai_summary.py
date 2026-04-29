"""Tests for the AI post-processing module.

The LLM call itself is mocked — the real OpenAI/Groq SDKs are
exercised in production. These tests cover the contract our wrapper
provides: input validation, JSON-mode parsing, schema normalisation
(the model occasionally returns slight shape drift), and the
exporter integration that puts the summary into DOCX/PDF.
"""

import io
import json
import os
import sys
from types import SimpleNamespace
from unittest.mock import MagicMock, patch

import pytest

PROJECT_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
sys.path.insert(0, PROJECT_ROOT)

from transcriber import ai_summary, batch, exporter
from transcriber.ai_summary import (
    SummaryError,
    _coerce_string_list,
    _normalise_summary_payload,
    summarize_transcript,
)


# ---------------------------------------------------------------------------
# Input validation
# ---------------------------------------------------------------------------

def test_empty_transcript_returns_empty_payload():
    """Empty input is a valid no-op — it would be a needless API call
    and the cleanest answer is "nothing to summarise"."""
    out = summarize_transcript("", provider="Groq (llama-3.3-70b)", api_key="x")
    assert out == {"summary": "", "topics": [], "action_items": []}

    out = summarize_transcript("   \n  ", provider="Groq (llama-3.3-70b)", api_key="x")
    assert out == {"summary": "", "topics": [], "action_items": []}


def test_unknown_provider_raises_summary_error():
    with pytest.raises(SummaryError, match="Unknown summary provider"):
        summarize_transcript("Hello world.", provider="Anthropic Claude", api_key="x")


def test_oversized_transcript_raises_with_helpful_message():
    """The soft cap protects users from a context-overflow failure
    that returns an opaque API error after a long retry storm."""
    huge = "word " * 80_000  # ~400k chars, past the cap
    with pytest.raises(SummaryError, match="too long"):
        summarize_transcript(huge, provider="Groq (llama-3.3-70b)", api_key="x")


# ---------------------------------------------------------------------------
# Payload normalisation
# ---------------------------------------------------------------------------

def test_normalise_summary_payload_passthrough():
    payload = {
        "summary": "A meeting about hiring.",
        "topics": ["Hiring", "Budget"],
        "action_items": ["Send offer letter."],
    }
    assert _normalise_summary_payload(payload) == payload


def test_normalise_summary_payload_strips_whitespace():
    out = _normalise_summary_payload({
        "summary": "  A meeting.  ",
        "topics": ["  Hiring  ", "Budget"],
        "action_items": ["  Send offer.  "],
    })
    assert out["summary"] == "A meeting."
    assert out["topics"] == ["Hiring", "Budget"]
    assert out["action_items"] == ["Send offer."]


def test_normalise_summary_payload_handles_missing_keys():
    """A model that returns only ``{"summary": "..."}`` shouldn't
    crash the renderer with KeyError."""
    out = _normalise_summary_payload({"summary": "Just a summary."})
    assert out["topics"] == []
    assert out["action_items"] == []


def test_normalise_summary_payload_rejects_non_dict():
    with pytest.raises(SummaryError, match="expected a JSON object"):
        _normalise_summary_payload(["not a dict"])


# ---------------------------------------------------------------------------
# String-list coercion (the model occasionally fights JSON mode)
# ---------------------------------------------------------------------------

def test_coerce_string_list_accepts_list():
    assert _coerce_string_list(["a", "b", ""]) == ["a", "b"]


def test_coerce_string_list_splits_string_on_newlines():
    """Models occasionally collapse a list into a multi-line string
    when JSON mode is unfamiliar with the requested schema."""
    raw = "- First item\n- Second item\n* Third item"
    assert _coerce_string_list(raw) == ["First item", "Second item", "Third item"]


def test_coerce_string_list_returns_empty_for_none_or_garbage():
    assert _coerce_string_list(None) == []
    assert _coerce_string_list(42) == []
    assert _coerce_string_list({"key": "value"}) == []


# ---------------------------------------------------------------------------
# End-to-end with mocked LLM
# ---------------------------------------------------------------------------

def _mock_chat_response(json_str: str):
    """Build a fake chat-completion response matching the OpenAI/Groq SDK shape."""
    return SimpleNamespace(
        choices=[SimpleNamespace(message=SimpleNamespace(content=json_str))],
    )


def test_happy_path_returns_normalised_payload(monkeypatch):
    """A well-formed JSON response from the model lands in the caller
    as a clean dict with all three fields — verifies the full path
    through the SDK call, JSON parse, and normalisation."""
    captured_kwargs = {}

    fake_client = MagicMock()
    fake_client.chat.completions.create.return_value = _mock_chat_response(
        json.dumps({
            "summary": "A discussion about Q3 hiring.",
            "topics": ["Hiring plan", "Budget approval"],
            "action_items": ["Maria will draft the offer letter by Friday."],
        })
    )

    def fake_factory(*, api_key):
        captured_kwargs["api_key"] = api_key
        return fake_client

    monkeypatch.setitem(
        ai_summary.SUMMARY_PROVIDERS,
        "Groq (llama-3.3-70b)",
        {**ai_summary.SUMMARY_PROVIDERS["Groq (llama-3.3-70b)"], "client_factory": fake_factory},
    )

    out = summarize_transcript(
        "Maria: We need to hire two engineers this quarter.",
        provider="Groq (llama-3.3-70b)",
        api_key="dummy-key",
    )

    assert out["summary"].startswith("A discussion")
    assert out["topics"] == ["Hiring plan", "Budget approval"]
    assert out["action_items"] == ["Maria will draft the offer letter by Friday."]
    assert captured_kwargs["api_key"] == "dummy-key"

    # The wrapper must request JSON mode — the system prompt alone
    # isn't reliable enough on smaller models.
    call_kwargs = fake_client.chat.completions.create.call_args.kwargs
    assert call_kwargs["response_format"] == {"type": "json_object"}


def test_invalid_json_response_raises_summary_error(monkeypatch):
    """A plain-prose answer (model ignored JSON mode) must produce a
    typed error the UI can render against — not a raw JSONDecodeError."""
    fake_client = MagicMock()
    fake_client.chat.completions.create.return_value = _mock_chat_response(
        "Sure! Here is your summary: ..."
    )
    monkeypatch.setitem(
        ai_summary.SUMMARY_PROVIDERS,
        "Groq (llama-3.3-70b)",
        {**ai_summary.SUMMARY_PROVIDERS["Groq (llama-3.3-70b)"], "client_factory": lambda **k: fake_client},
    )

    with pytest.raises(SummaryError, match="invalid JSON"):
        summarize_transcript(
            "Hello.",
            provider="Groq (llama-3.3-70b)",
            api_key="x",
        )


def test_empty_response_raises_summary_error(monkeypatch):
    fake_client = MagicMock()
    fake_client.chat.completions.create.return_value = _mock_chat_response("")
    monkeypatch.setitem(
        ai_summary.SUMMARY_PROVIDERS,
        "Groq (llama-3.3-70b)",
        {**ai_summary.SUMMARY_PROVIDERS["Groq (llama-3.3-70b)"], "client_factory": lambda **k: fake_client},
    )

    with pytest.raises(SummaryError, match="empty response"):
        summarize_transcript("Hello.", provider="Groq (llama-3.3-70b)", api_key="x")


# ---------------------------------------------------------------------------
# Exporter integration: summary lands in DOCX + PDF
# ---------------------------------------------------------------------------

_SAMPLE_SUMMARY = {
    "summary": "Discussion about Q3 hiring and budget.",
    "topics": ["Hiring plan", "Budget approval", "Timeline"],
    "action_items": ["Maria sends offer letter.", "John reviews budget."],
}


def test_export_docx_includes_summary_text():
    """All summary fields must round-trip into the DOCX body so a
    reader without the editor still sees the metadata."""
    blob = exporter.export_docx(
        "**Speaker 0:**\nHello.",
        title="Test",
        summary=_SAMPLE_SUMMARY,
    )
    from docx import Document

    doc = Document(io.BytesIO(blob))
    rendered = "\n".join(p.text for p in doc.paragraphs)
    assert "AI Summary" in rendered
    assert "Discussion about Q3 hiring" in rendered
    assert "Hiring plan" in rendered
    assert "Maria sends offer letter." in rendered


def test_export_docx_without_summary_unchanged():
    """The summary section must be absent when no summary is passed —
    we can't have empty headings polluting an export from a user who
    didn't enable the feature."""
    blob = exporter.export_docx("Plain transcript.", title="Test")
    from docx import Document

    doc = Document(io.BytesIO(blob))
    rendered = "\n".join(p.text for p in doc.paragraphs)
    assert "AI Summary" not in rendered


def test_export_pdf_with_summary_produces_valid_pdf():
    """fpdf2 zlib-compresses content streams, so we can't grep for the
    summary text directly without parsing the PDF. The narrower
    invariant: passing a summary must produce a valid, longer PDF
    than the same call without one — proof the summary section
    actually emitted bytes into the stream rather than being
    silently dropped."""
    base = exporter.export_pdf("**Speaker 0:**\nHello.", title="Test")
    with_summary = exporter.export_pdf(
        "**Speaker 0:**\nHello.",
        title="Test",
        summary=_SAMPLE_SUMMARY,
    )
    assert base.startswith(b"%PDF")
    assert with_summary.startswith(b"%PDF")
    assert len(with_summary) > len(base), (
        "summary section should add bytes to the PDF stream"
    )


def test_batch_zip_includes_per_file_summary_in_docx():
    """End-to-end: a batch entry that carries a summary must produce
    a DOCX inside the zip that contains the summary text."""
    import zipfile
    blob = batch.build_zip([{
        "filename": "meeting.mp4",
        "text": "**Speaker 0:**\nHello.",
        "summary": _SAMPLE_SUMMARY,
    }])
    with zipfile.ZipFile(io.BytesIO(blob)) as zf:
        docx_bytes = zf.read("meeting/transcription.docx")

    from docx import Document
    doc = Document(io.BytesIO(docx_bytes))
    rendered = "\n".join(p.text for p in doc.paragraphs)
    assert "AI Summary" in rendered
    assert "Hiring plan" in rendered
