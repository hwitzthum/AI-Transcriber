"""Export transcribed text to DOCX and PDF formats."""

import io
import datetime
import os
import re
import sys
from typing import Optional, Tuple, List

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_COLOR_INDEX
from fpdf import FPDF


# Inline markers the transcript pipeline can emit. Both ``_filler_`` and
# ``~~low-confidence~~`` are rendered into the formatted segment list
# below; ``_strip_inline_markers`` peels them off when an export path
# can't represent the styling (e.g. PDF body text).
_FILLER_PATTERN = re.compile(r"_([^_]+)_")
_LOW_CONFIDENCE_PATTERN = re.compile(r"~~([^~]+)~~")
# Combined pattern: matches ``~~text~~`` (group 1) OR ``_text_``
# (group 2). Used by the segment splitter so a single pass handles
# both marker shapes in source order.
_INLINE_MARKER_PATTERN = re.compile(r"~~([^~\n]+?)~~|_([^_\n]+?)_")

# Leading ``[HH:MM:SS]`` marker stripped off before we look for the
# speaker label. Same shape the timestamped-transcript pipeline emits.
_LEADING_TIMESTAMP = re.compile(r"^\s*\[(\d{2}:\d{2}:\d{2})\]\s*")


# Module-level font cache: font_path for the Unicode font
_cached_font_path: Optional[str] = None


def _find_unicode_font() -> Optional[str]:
    """
    Find a Unicode-capable font for PDF generation.

    Prioritizes fonts with good Unicode support (French œ, German ü, etc.)

    Returns:
        Full path to font file if found, None otherwise.
    """
    # Font candidates in priority order (best Unicode support first)
    font_candidates = []

    if sys.platform == "darwin":  # macOS
        font_candidates = [
            # Arial Unicode has excellent Unicode support
            "/System/Library/Fonts/Supplemental/Arial Unicode.ttf",
            "/Library/Fonts/Arial Unicode.ttf",
            # Standard Arial in Supplemental folder
            "/System/Library/Fonts/Supplemental/Arial.ttf",
            # DejaVu if installed via Homebrew
            "/opt/homebrew/share/fonts/dejavu/DejaVuSans.ttf",
            "/usr/local/share/fonts/dejavu/DejaVuSans.ttf",
        ]
    elif sys.platform == "linux":
        font_candidates = [
            "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
            "/usr/share/fonts/dejavu/DejaVuSans.ttf",
            "/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf",
            "/usr/share/fonts/truetype/freefont/FreeSans.ttf",
        ]
    elif sys.platform == "win32":
        font_candidates = [
            r"C:\Windows\Fonts\arial.ttf",
            r"C:\Windows\Fonts\ArialUni.ttf",
            r"C:\Windows\Fonts\calibri.ttf",
        ]

    # Check each candidate
    for font_path in font_candidates:
        if os.path.isfile(font_path):
            return font_path

    return None


def _create_pdf_with_unicode_font() -> Tuple[FPDF, str]:
    """
    Create an FPDF instance with a Unicode font registered.

    Returns:
        Tuple of (FPDF instance, font family name to use)
    """
    global _cached_font_path

    pdf = FPDF()

    # Try to find and register a Unicode font
    if _cached_font_path is None:
        _cached_font_path = _find_unicode_font()

    if _cached_font_path and os.path.isfile(_cached_font_path):
        try:
            pdf.add_font("UniFont", "", _cached_font_path, uni=True)
            return pdf, "UniFont"
        except Exception:
            pass

    # Fallback to Helvetica (limited charset)
    return pdf, "Helvetica"


def _parse_speaker_block(block: str) -> Tuple[Optional[str], Optional[str], str]:
    """
    Parse a text block to extract optional timestamp, speaker label, and content.

    Handles four shapes the transcript pipeline can produce:

    1. ``**Speaker X:**\\n<content>``                    — diarized, no timestamp
    2. ``[HH:MM:SS] **Speaker X:**\\n<content>``         — diarized + timestamped
    3. ``[HH:MM:SS]\\n<content>``                        — non-diarized + timestamped
    4. ``<content>``                                     — plain prose

    Args:
        block: A text block from the transcript (one of the shapes above).

    Returns:
        ``(timestamp, speaker_label, content)`` where each of the first
        two may be ``None``. Timestamp is returned as ``HH:MM:SS`` (no
        brackets); speaker label is returned without the ``**`` markers.
    """
    stripped = block.strip()
    if not stripped:
        return None, None, ""

    # Peel off a leading [HH:MM:SS] marker first so the speaker check
    # below sees the same shape regardless of whether timestamps are on.
    timestamp: Optional[str] = None
    ts_match = _LEADING_TIMESTAMP.match(stripped)
    if ts_match:
        timestamp = ts_match.group(1)
        stripped = stripped[ts_match.end():].lstrip()

    lines = stripped.split("\n", 1)
    first_line = lines[0].strip()
    rest = lines[1].strip() if len(lines) > 1 else ""

    match = re.match(r"^\*\*(.+?):\*\*$", first_line)
    if match:
        speaker_label = f"{match.group(1)}:"
        return timestamp, speaker_label, rest

    # No speaker label — the whole timestamp-stripped block is content.
    return timestamp, None, stripped


def _parse_text_segments(text: str) -> List[Tuple[str, Optional[str]]]:
    """
    Parse text into ``(text, kind)`` segments where ``kind`` is one of
    ``"filler"``, ``"low_confidence"``, or ``None`` (unstyled).

    Iterating over the combined pattern (rather than running two passes,
    one per marker type) preserves source order, so when the renderer
    walks the result a span like ``"_um_ ~~probably~~"`` produces an
    italic-gray run followed by an amber-highlighted run in that order.
    """
    segments: List[Tuple[str, Optional[str]]] = []
    last_end = 0

    for match in _INLINE_MARKER_PATTERN.finditer(text):
        if match.start() > last_end:
            segments.append((text[last_end:match.start()], None))

        if match.group(1) is not None:
            segments.append((match.group(1), "low_confidence"))
        else:
            segments.append((match.group(2), "filler"))

        last_end = match.end()

    if last_end < len(text):
        segments.append((text[last_end:], None))

    return segments if segments else [(text, None)]


def _add_formatted_text_to_paragraph(para, text: str) -> None:
    """
    Add text to a DOCX paragraph with filler / low-confidence styling.

    Filler words render as italic gray (matches the editor preview).
    Low-confidence words get a yellow text highlight so a reviewer can
    spot what to verify without reading the whole transcript.
    """
    segments = _parse_text_segments(text)

    for segment_text, kind in segments:
        run = para.add_run(segment_text)
        if kind == "filler":
            run.italic = True
            run.font.color.rgb = RGBColor(0x9c, 0xa3, 0xaf)
        elif kind == "low_confidence":
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def _strip_inline_markers(text: str) -> str:
    """
    Remove ``_filler_`` AND ``~~low-confidence~~`` markers from a string,
    keeping the words. Used by the PDF export, where the body-text
    renderer can't apply per-run styling without a major restructure —
    so the cleaner option is to drop the markers and produce a plain
    transcript that's still readable. The inline visual cues are
    available in the editor preview and the DOCX export.
    """
    cleaned = _FILLER_PATTERN.sub(r"\1", text)
    cleaned = _LOW_CONFIDENCE_PATTERN.sub(r"\1", cleaned)
    return cleaned


def export_docx(text: str, title: str = "Transcription") -> bytes:
    """
    Export text to a DOCX document with formatted speaker labels.

    Speaker labels (e.g., **Speaker 0:**) are rendered bold on their own line,
    followed by the speaker's text.

    Args:
        text: The transcribed text content.
        title: Document title.

    Returns:
        DOCX file content as bytes.
    """
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # Title
    heading = doc.add_heading(title, level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Metadata
    meta_para = doc.add_paragraph()
    meta_run = meta_para.add_run(
        f"Generated on {datetime.datetime.now().strftime('%Y-%m-%d %H:%M')}"
    )
    meta_run.font.size = Pt(9)
    meta_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    meta_run.font.italic = True

    # Separator
    doc.add_paragraph("─" * 60)

    # Content — split by double newlines into paragraphs
    paragraphs = text.split("\n\n")
    for para_text in paragraphs:
        para_text = para_text.strip()
        if not para_text:
            continue

        timestamp, speaker_label, content = _parse_speaker_block(para_text)

        if speaker_label:
            # Speaker label paragraph: optional timestamp prefix in muted
            # gray, then the speaker name in bold blue. One paragraph so
            # the eye reads "[time] · Speaker:" as a single header line.
            speaker_para = doc.add_paragraph()
            if timestamp:
                ts_run = speaker_para.add_run(f"[{timestamp}]  ")
                ts_run.font.size = Pt(9)
                ts_run.font.color.rgb = RGBColor(0x9c, 0xa3, 0xaf)
                ts_run.font.name = "Consolas"
            speaker_run = speaker_para.add_run(speaker_label)
            speaker_run.bold = True
            speaker_run.font.size = Pt(11)
            speaker_run.font.color.rgb = RGBColor(0x1a, 0x56, 0xdb)
            speaker_para.paragraph_format.space_after = Pt(2)

            # Add content paragraph with filler word formatting
            if content:
                content_para = doc.add_paragraph()
                _add_formatted_text_to_paragraph(content_para, content)
                content_para.paragraph_format.space_after = Pt(12)
                content_para.paragraph_format.left_indent = Inches(0.25)
        else:
            # Plain paragraph (possibly preceded by a timestamp). The
            # timestamp goes on its own line in muted gray so the body
            # text isn't visually broken up.
            if timestamp:
                ts_para = doc.add_paragraph()
                ts_run = ts_para.add_run(f"[{timestamp}]")
                ts_run.font.size = Pt(9)
                ts_run.font.color.rgb = RGBColor(0x9c, 0xa3, 0xaf)
                ts_run.font.name = "Consolas"
                ts_para.paragraph_format.space_after = Pt(2)

            para = doc.add_paragraph()
            _add_formatted_text_to_paragraph(para, content)
            para.paragraph_format.space_after = Pt(8)

    # Save to bytes buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def export_pdf(text: str, title: str = "Transcription") -> bytes:
    """
    Export text to a PDF document with formatted speaker labels.

    Speaker labels (e.g., **Speaker 0:**) are rendered in bold blue on their own line,
    followed by the speaker's text with slight indentation.

    Args:
        text: The transcribed text content.
        title: Document title.

    Returns:
        PDF file content as bytes.

    Raises:
        ValueError: If text is empty.
    """
    # Handle empty text input
    if not text or not text.strip():
        raise ValueError("Cannot export empty text to PDF")

    # Create PDF with Unicode font registered on THIS instance
    pdf, body_font = _create_pdf_with_unicode_font()
    pdf.set_auto_page_break(auto=True, margin=25)
    pdf.add_page()

    # Title
    pdf.set_font(body_font, "", 18)
    pdf.set_text_color(33, 33, 33)
    pdf.cell(0, 15, title, new_x="LMARGIN", new_y="NEXT")

    # Metadata
    pdf.set_font(body_font, "", 9)
    pdf.set_text_color(153, 153, 153)
    pdf.cell(
        0, 8,
        f"Generated on {datetime.datetime.now().strftime('%Y-%m-%d %H:%M')}",
        new_x="LMARGIN", new_y="NEXT",
    )

    # Separator line
    pdf.set_draw_color(200, 200, 200)
    pdf.line(10, pdf.get_y() + 2, 200, pdf.get_y() + 2)
    pdf.ln(8)

    # Content - process paragraphs with speaker label formatting
    paragraphs = text.split("\n\n")
    for para_text in paragraphs:
        para_text = para_text.strip()
        if not para_text:
            continue

        timestamp, speaker_label, content = _parse_speaker_block(para_text)

        if speaker_label:
            # Optional [HH:MM:SS] prefix in muted gray on its own line
            # above the speaker label, so the speaker name still stands
            # out as the visual anchor for the block.
            if timestamp:
                pdf.set_font(body_font, "", 9)
                pdf.set_text_color(156, 163, 175)
                pdf.cell(0, 5, f"[{timestamp}]", new_x="LMARGIN", new_y="NEXT")

            pdf.set_font(body_font, "", 12)
            pdf.set_text_color(26, 86, 219)
            pdf.cell(0, 7, speaker_label, new_x="LMARGIN", new_y="NEXT")

            # Content: normal text, slightly indented (strip filler markers)
            if content:
                pdf.set_font(body_font, "", 11)
                pdf.set_text_color(51, 51, 51)
                pdf.set_x(pdf.l_margin + 5)  # Slight indent
                pdf.multi_cell(0, 6, _strip_inline_markers(content))
            pdf.ln(6)
        else:
            # Plain paragraph; same muted-timestamp treatment as above.
            if timestamp:
                pdf.set_font(body_font, "", 9)
                pdf.set_text_color(156, 163, 175)
                pdf.cell(0, 5, f"[{timestamp}]", new_x="LMARGIN", new_y="NEXT")

            pdf.set_font(body_font, "", 11)
            pdf.set_text_color(51, 51, 51)
            pdf.multi_cell(0, 6, _strip_inline_markers(content))
            pdf.ln(4)

    # Output as bytes - handle bytes, bytearray, and string return types from fpdf2
    output = pdf.output()
    if isinstance(output, (bytes, bytearray)):
        return bytes(output)
    return output.encode("utf-8")